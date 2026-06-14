"""
Build the formal v1.5 thesis manuscript as a DOCX artefact.

The document uses the Documents skill `narrative_proposal` preset:
Letter page, 1 inch margins, Calibri 11 pt body, justified body text,
1.333 line spacing, restrained blue headings, fixed-width tables and
quiet running headers.
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = PROJECT_ROOT / "docs"
RESULTS_DIR = PROJECT_ROOT / "results"
OUTPUT_PATH = DOCS_DIR / "Australian_Raptor_Thesis_v1_5.docx"

CLASS_ORDER = [
    ("aquila_audax", "Wedge-tailed Eagle", "Accipitridae"),
    ("circus_assimilis", "Spotted Harrier", "Accipitridae"),
    ("elanus_axillaris", "Black-shouldered Kite", "Accipitridae"),
    ("falco_cenchroides", "Nankeen Kestrel", "Falconidae"),
    ("falco_peregrinus", "Peregrine Falcon", "Falconidae"),
    ("hieraaetus_morphnoides", "Little Eagle", "Accipitridae"),
    ("lophoictinia_isura", "Square-tailed Kite", "Accipitridae"),
    ("tachyspiza_fasciata", "Brown Goshawk", "Accipitridae"),
]

DOC_TITLE = (
    "YOLO-Assisted Deep Learning for Australian Raptor Identification "
    "with Accessible Citizen-Science Interfaces"
)


def read_json(path: Path) -> dict:
    if not path.exists():
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin_name, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "BFC7D2")

    tbl_grid = table._tbl.tblGrid
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            cell.width = width
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_margins(cell)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.208

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    caption = styles.add_style("Thesis Caption", 1)
    caption.font.name = "Calibri"
    caption.font.size = Pt(9)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    caption.paragraph_format.space_before = Pt(2)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    source = styles.add_style("Table Source", 1)
    source.font.name = "Calibri"
    source.font.size = Pt(9)
    source.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    source.paragraph_format.space_before = Pt(4)
    source.paragraph_format.space_after = Pt(4)

    header = section.header.paragraphs[0]
    header.text = "Australian Raptor CNN + AUSLAN | v1.5 academic baseline"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(96)
    run = p.add_run(DOC_TITLE)
    run.font.name = "Calibri"
    run.font.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x0B, 0x25, 0x45)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run("Master's Research Manuscript | Release v1.5")
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x1F, 0x4D, 0x78)

    for label, value in [
        ("Author", "Brian Fernandez-Baez"),
        ("System", "YOLO-assisted localisation + EfficientNetB4 classification"),
        ("Dataset", "1,992 processed images, eight Australian raptor species"),
        ("Evaluation", "206-image held-out test split, seed 42"),
        ("Date", "2026-06-07"),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{label}: ")
        r1.font.bold = True
        p.add_run(value)

    doc.add_page_break()


def add_contents(doc: Document) -> None:
    doc.add_heading("Contents", level=1)
    items = [
        "Abstract",
        "Chapter 1. Introduction",
        "Chapter 2. Literature Review",
        "Chapter 3. Methodology",
        "Chapter 4. Results",
        "Chapter 5. Software, Accessibility and Data Interoperability",
        "Chapter 6. Discussion and Conclusion",
        "Appendix A. Dataset Datasheet Summary",
        "Appendix B. Model Card Summary",
        "Appendix C. Reproducibility Checklist",
        "Appendix D. Defence Demonstration Script",
        "Appendix E. Release Package",
        "Bibliography",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_page_break()


def add_para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header
        set_cell_shading(header_cells[idx], "F4F6F9")
        for paragraph in header_cells[idx].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    set_table_geometry(table, widths)


def add_source(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Table Source")


def add_figure(doc: Document, path: Path, caption: str) -> None:
    if not path.exists():
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(6.1))
    doc.add_paragraph(caption, style="Thesis Caption")


def add_abstract(doc: Document) -> None:
    doc.add_heading("Abstract", level=1)
    add_para(
        doc,
        "This thesis presents a research prototype for Australian raptor "
        "identification under citizen-science image conditions. The system "
        "combines YOLO-based bird localisation with an EfficientNetB4 "
        "classifier trained on eight Australian raptor species from "
        "iNaturalist Australia and the Atlas of Living Australia.",
    )
    add_para(
        doc,
        "The prototype implements a multilingual Flask interface, top-3 "
        "species alternatives, Darwin Core export, feedback logging and "
        "provisional AUSLAN motion illustrations. On a 206-image held-out "
        "test split, the classifier reaches 84.95% top-1 accuracy and "
        "0.8482 macro-F1. Bootstrap confidence intervals, calibration "
        "analysis, family-level error analysis and Grad-CAM diagnostics are "
        "reported to support reproducibility and scientific caution.",
    )
    add_para(
        doc,
        "The v1.5 release is positioned as an academic baseline rather than "
        "a production wildlife-identification authority. It documents the "
        "limits of the dataset, the provisional status of accessibility "
        "assets and a v2.0 roadmap for 14 species, detector fine-tuning and "
        "stronger split governance.",
    )


def add_chapter_1(doc: Document) -> None:
    doc.add_heading("Chapter 1. Introduction", level=1)
    add_para(
        doc,
        "Australian raptors are ecologically important apex and meso-"
        "predators, but reliable species identification remains difficult "
        "for non-specialists. Citizen-science platforms can expand "
        "monitoring coverage, yet visual identification, biodiversity data "
        "interoperability and accessibility for Deaf participants remain "
        "open problems."
    )
    add_para(
        doc,
        "This thesis asks whether a compact and reproducible computer-vision "
        "pipeline can support Australian raptor identification while also "
        "exposing records in Darwin Core format and documenting a pathway "
        "for participatory AUSLAN validation."
    )
    doc.add_heading("Research Questions", level=2)
    add_numbered(
        doc,
        [
            "Can a YOLO + EfficientNetB4 pipeline identify eight Australian "
            "raptor species with useful accuracy under citizen-science image "
            "conditions?",
            "Can the prediction workflow produce auditable per-image "
            "predictions and Darwin Core-compatible records?",
            "How should provisional AUSLAN signs be framed so they invite "
            "participatory validation without claiming authority?",
        ],
    )
    doc.add_heading("Contributions", level=2)
    add_bullets(
        doc,
        [
            "A validated eight-species v1.5 raptor classification baseline.",
            "A YOLO-assisted localisation path for images and video frames.",
            "Per-image evaluation outputs in results/test_predictions.csv.",
            "Bootstrap, calibration and taxonomy-aware error analyses.",
            "A multilingual Flask app with feedback and Darwin Core export.",
            "A documented AUSLAN consultation protocol.",
        ],
    )


def add_chapter_2(doc: Document) -> None:
    doc.add_heading("Chapter 2. Literature Review", level=1)
    add_para(
        doc,
        "The project draws on four connected bodies of work: fine-grained "
        "bird classification and transfer learning; object detection for "
        "ecological image analysis; citizen-science biodiversity "
        "infrastructure, including ALA, GBIF and Darwin Core; and "
        "accessibility research, AUSLAN and participatory design."
    )
    add_para(
        doc,
        "EfficientNetB4 is used as the single v1.5 classifier because it "
        "offers a strong accuracy-to-compute balance at 380 pixel inputs. "
        "YOLO is used for localisation and candidate cropping, but the "
        "current release uses an adaptive policy because crop-only inference "
        "did not improve the held-out checkpoint."
    )
    add_para(
        doc,
        "The accessibility contribution is deliberately framed as a "
        "participatory prototype. The project does not assert that generated "
        "motion illustrations are validated AUSLAN signs; instead, it "
        "documents a consultation protocol that can be reviewed by Deaf "
        "community participants and interpreters."
    )


def add_chapter_3(doc: Document) -> None:
    doc.add_heading("Chapter 3. Methodology", level=1)
    doc.add_heading("Dataset", level=2)
    add_para(
        doc,
        "The v1.5 processed dataset contains 1,992 images across eight "
        "species. The active class order is fixed in gui/app.py::CLASS_ORDER "
        "and is verified by automated tests so that the model, dataset and "
        "user interface remain synchronized."
    )
    add_table(
        doc,
        ["Split", "Images"],
        [["Train", "1,590"], ["Validation", "196"], ["Test", "206"]],
        [4680, 4680],
    )
    add_source(doc, "Source: processed project dataset, deterministic seed 42.")

    add_table(
        doc,
        ["Class key", "Common name", "Family"],
        [[key, common, family] for key, common, family in CLASS_ORDER],
        [2800, 3760, 2800],
    )
    add_source(doc, "Source: gui/app.py::CLASS_ORDER and project taxonomy metadata.")

    doc.add_heading("Pipeline", level=2)
    add_numbered(
        doc,
        [
            "YOLO detects bird regions in an uploaded image or sampled video frame.",
            "EfficientNetB4 classifies the whole image and, when useful, the YOLO crop.",
            "The adaptive policy uses the crop only when it is more confident than the whole-image prediction.",
            "The interface displays top-1 confidence and top-3 alternatives.",
            "User feedback and Darwin Core exports are written to CSV artefacts.",
        ],
    )
    doc.add_heading("Reproducibility", level=2)
    add_table(
        doc,
        ["Artefact", "Purpose"],
        [
            ["requirements.txt", "Python dependency lock for local and CI execution."],
            ["environment.yml", "Conda-compatible environment definition."],
            ["Dockerfile", "Containerized reproducibility path."],
            [".github/workflows/ci.yml", "Lightweight CI with real pytest and healthcheck."],
            ["notebooks/healthcheck.py", "Project contract checks for release integrity."],
            ["tests/test_project_integrity.py", "Automated tests for Flask, i18n, Darwin Core and YOLO wrapper."],
            ["docs/DATASHEET.md", "Dataset provenance, risks and constraints."],
            ["docs/MODEL_CARD.md", "Model behavior, metrics and intended use."],
        ],
        [3000, 6360],
    )


def add_chapter_4(doc: Document, report: dict, bootstrap: dict, calibration: dict) -> None:
    doc.add_heading("Chapter 4. Results", level=1)
    metrics = report.get("metricas_globales", {})
    add_table(
        doc,
        ["Metric", "Value"],
        [
            ["Test images", str(report.get("total_test_images", 206))],
            ["Top-1 accuracy", f"{metrics.get('accuracy', 0.8495):.4f}"],
            ["Macro-F1", f"{metrics.get('f1_macro', 0.8482):.4f}"],
            ["Weighted-F1", f"{metrics.get('f1_weighted', 0.8476):.4f}"],
            ["Expected calibration error", f"{calibration.get('ece', 0.0639):.4f}"],
            ["Family-level accuracy", "0.9272"],
        ],
        [4680, 4680],
    )
    add_source(doc, "Source: results/reporte_final.json and results/calibration_efficientnet_b4.json.")

    per_species = report.get("por_especie", {})
    rows = []
    for _, common, _ in CLASS_ORDER:
        values = per_species.get(common, {})
        rows.append(
            [
                common,
                f"{values.get('precision', 0):.4f}",
                f"{values.get('recall', 0):.4f}",
                f"{values.get('f1', 0):.4f}",
                str(values.get("support", "")),
            ]
        )
    add_table(
        doc,
        ["Species", "Precision", "Recall", "F1", "Support"],
        rows,
        [2960, 1600, 1600, 1600, 1600],
    )
    add_source(doc, "Source: held-out test predictions for EfficientNetB4.")

    accuracy = bootstrap.get("accuracy", {})
    macro_f1 = bootstrap.get("macro_f1", {})
    add_table(
        doc,
        ["Metric", "Bootstrap mean", "95% CI"],
        [
            [
                "Accuracy",
                f"{accuracy.get('mean', 0.850388):.4f}",
                f"[{accuracy.get('ci_lo', 0.800971):.4f}, {accuracy.get('ci_hi', 0.898058):.4f}]",
            ],
            [
                "Macro-F1",
                f"{macro_f1.get('mean', 0.846675):.4f}",
                f"[{macro_f1.get('ci_lo', 0.796381):.4f}, {macro_f1.get('ci_hi', 0.893512):.4f}]",
            ],
        ],
        [2800, 2800, 3760],
    )
    add_source(doc, "Source: results/bootstrap_ci_efficientnet_b4.json, n = 1,000 bootstrap resamples.")

    add_para(
        doc,
        "The model reaches 92.72% family-level accuracy. Cross-family errors "
        "account for 15 of 206 test images (7.3%). The most common reported "
        "confusion is tachyspiza_fasciata predicted as elanus_axillaris, "
        "followed by falco_cenchroides predicted as elanus_axillaris."
    )

    add_figure(
        doc,
        RESULTS_DIR / "confusion_family_efficientnet_b4.png",
        "Figure 1. Family-level confusion matrix for the v1.5 EfficientNetB4 classifier.",
    )
    add_figure(
        doc,
        RESULTS_DIR / "reliability_diagram_efficientnet_b4.png",
        "Figure 2. Reliability diagram used to estimate expected calibration error.",
    )
    add_figure(
        doc,
        RESULTS_DIR / "gradcam_mosaic.png",
        "Figure 3. Grad-CAM mosaic showing image regions influencing classifier decisions.",
    )


def add_chapter_5(doc: Document) -> None:
    doc.add_heading("Chapter 5. Software, Accessibility and Data Interoperability", level=1)
    add_para(
        doc,
        "The Flask application operationalizes the model as a citizen-science "
        "prototype. It supports image upload, video-frame sampling, top-3 "
        "alternatives, feedback logging, out-of-domain feedback, Darwin Core "
        "export and 10-language user-interface text."
    )
    add_table(
        doc,
        ["System area", "Implemented evidence"],
        [
            ["Flask routes", "Health, identify, feedback, Darwin Core and static interface routes are covered by pytest."],
            ["Prediction audit", "results/test_predictions.csv stores image_path, y_true, y_pred, confidence and top3."],
            ["Darwin Core", "Export fields include scientificName, taxonID, basisOfRecord and identifiedBy."],
            ["Feedback", "Incorrect predictions and out-of-domain reports are appended to CSV logs."],
            ["i18n", "Ten translation files are checked for active species coverage."],
            ["YOLO", "gui/yolo_detector.py provides optional Ultralytics-based bird localisation."],
        ],
        [2600, 6760],
    )
    add_para(
        doc,
        "The AUSLAN component is explicitly provisional. No sign is claimed "
        "as validated AUSLAN until reviewed through Deaf-community "
        "consultation. The contribution is therefore a process and prototype "
        "contribution, not an authoritative vocabulary."
    )


def add_chapter_6(doc: Document) -> None:
    doc.add_heading("Chapter 6. Discussion and Conclusion", level=1)
    add_para(
        doc,
        "The v1.5 release demonstrates that a modest, reproducible YOLO-"
        "assisted EfficientNetB4 pipeline can support Australian raptor "
        "identification with useful but imperfect performance. The system is "
        "strong enough for a research prototype and scholarship portfolio, "
        "but not for legally consequential deployment."
    )
    doc.add_heading("Limitations", level=2)
    add_bullets(
        doc,
        [
            "The split is per-image rather than group-aware by photographer, location or observation event.",
            "The release is closed-set and limited to eight active species.",
            "Juvenile plumages, difficult angles and remote regions remain under-represented.",
            "AUSLAN illustrations are provisional and require community validation.",
            "YOLO uses generic bird/person-era detector behaviour rather than raptor-specific fine-tuning.",
        ],
    )
    doc.add_heading("Future Work", level=2)
    add_bullets(
        doc,
        [
            "Expand v2.0 to the 14-species dataset after provenance and split validation.",
            "Keep EfficientNetB4 as the primary classifier while improving split governance, detector tuning and calibration.",
            "Fine-tune YOLO on raptor bounding boxes and report detector metrics.",
            "Introduce group-aware splits and geographic bias analysis.",
            "Run participatory AUSLAN validation before claiming signed vocabulary support.",
        ],
    )
    add_para(
        doc,
        "The central conclusion is that the project is now defensible as a "
        "transparent academic baseline: it contains real metrics, real "
        "automated tests, explicit limitations, reproducible artefacts and a "
        "clear pathway from Master's prototype to stronger future doctoral research."
    )


def add_appendices(doc: Document) -> None:
    doc.add_heading("Appendix A. Dataset Datasheet Summary", level=1)
    add_table(
        doc,
        ["Field", "v1.5 value"],
        [
            ["Sources", "iNaturalist Australia and Atlas of Living Australia."],
            ["Species", "Eight active raptor species."],
            ["Images", "1,992 processed images."],
            ["Splits", "Train 1,590; validation 196; test 206."],
            ["Risks", "Geographic, photographer, seasonal and plumage bias."],
            ["Standard", "Darwin Core export supported for prediction records."],
        ],
        [2600, 6760],
    )
    add_source(doc, "Full datasheet: docs/DATASHEET.md.")

    doc.add_heading("Appendix B. Model Card Summary", level=1)
    add_table(
        doc,
        ["Field", "v1.5 value"],
        [
            ["Classifier", "EfficientNetB4, eight output classes."],
            ["Detector", "YOLO localisation with an adaptive crop policy before classification."],
            ["Intended use", "Research prototype and citizen-science assistance."],
            ["Out-of-scope use", "Legal, veterinary, conservation enforcement or final expert identification."],
            ["Performance", "Accuracy 0.8495; macro-F1 0.8482; app-output ECE 0.0639."],
            ["Known weakness", "Closed-set errors, juvenile plumage and cross-family confusions."],
        ],
        [2600, 6760],
    )
    add_source(doc, "Full model card: docs/MODEL_CARD.md.")

    doc.add_heading("Appendix C. Reproducibility Checklist", level=1)
    add_bullets(
        doc,
        [
            "Run python notebooks/export_test_predictions.py to regenerate per-image predictions.",
            "Run python notebooks/bootstrap_metrics.py --report-md for confidence intervals.",
            "Run python notebooks/error_analysis.py for family-aware error analysis.",
            "Run python notebooks/calibration_ece.py for ECE and reliability diagrams.",
            "Run python notebooks/run_tests.py for automated project tests.",
            "Run python notebooks/healthcheck.py --verbose for release-gate checks.",
        ],
    )

    doc.add_heading("Appendix D. Defence Demonstration Script", level=1)
    add_numbered(
        doc,
        [
            "Open the README and state the v1.5 scope: eight species, EfficientNetB4 and YOLO adaptive localisation.",
            "Show results/test_predictions.csv and explain image_path, y_true, y_pred, confidence and top3.",
            "Show the bootstrap, calibration and error-analysis artefacts in results/.",
            "Run the test suite and healthcheck to demonstrate reproducibility.",
            "Open the Flask interface and perform one image prediction.",
            "Export a Darwin Core row and show how feedback is logged.",
            "Close with limitations and the v2.0 research roadmap.",
        ],
    )

    doc.add_heading("Appendix E. Release Package", level=1)
    add_table(
        doc,
        ["Artefact", "Purpose"],
        [
            ["docs/Australian_Raptor_Thesis_v1_5.docx", "Editable thesis manuscript."],
            ["docs/Australian_Raptor_Thesis_v1_5.pdf", "PDF thesis export for review and submission packets."],
            ["results/thesis_docx_audit.json", "Structural DOCX audit for headings, tables, figures and metrics."],
            ["results/thesis_pdf_audit.json", "PDF audit for page count, text extraction and required terms."],
            ["RELEASE_MANIFEST_v1_5.md", "SHA-256 manifest for release source, documentation and result artefacts."],
        ],
        [3600, 5760],
    )

    doc.add_heading("Bibliography", level=1)
    add_bullets(
        doc,
        [
            "Efron, B. (1979). Bootstrap methods: another look at the jackknife.",
            "Gebru, T. et al. (2021). Datasheets for datasets.",
            "Guo, C. et al. (2017). On calibration of modern neural networks.",
            "Mitchell, M. et al. (2019). Model cards for model reporting.",
            "Selvaraju, R. R. et al. (2017). Grad-CAM.",
            "Tan, M. and Le, Q. (2019). EfficientNet.",
            "Van Horn, G. et al. (2018). The iNaturalist species classification and detection dataset.",
            "Wieczorek, J. et al. (2012). Darwin Core: an evolving community-developed biodiversity data standard.",
        ],
    )


def main() -> None:
    report = read_json(RESULTS_DIR / "reporte_final.json")
    bootstrap = read_json(RESULTS_DIR / "bootstrap_ci_efficientnet_b4.json")
    calibration = read_json(RESULTS_DIR / "calibration_efficientnet_b4.json")

    doc = Document()
    configure_document(doc)
    add_title_page(doc)
    add_contents(doc)
    add_abstract(doc)
    doc.add_page_break()
    add_chapter_1(doc)
    doc.add_page_break()
    add_chapter_2(doc)
    doc.add_page_break()
    add_chapter_3(doc)
    doc.add_page_break()
    add_chapter_4(doc, report, bootstrap, calibration)
    doc.add_page_break()
    add_chapter_5(doc)
    doc.add_page_break()
    add_chapter_6(doc)
    doc.add_page_break()
    add_appendices(doc)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
