"""
Structural audit for the generated thesis DOCX.

This is a fallback QA path when LibreOffice rendering is unavailable. It does
not replace visual page review, but it verifies that the manuscript contains
the required academic sections, tables, embedded figures and style markers.
"""

from __future__ import annotations

import json
import zipfile
from pathlib import Path

from docx import Document


PROJECT_ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = PROJECT_ROOT / "docs" / "Australian_Raptor_Thesis_v1_5.docx"
RESULTS_PATH = PROJECT_ROOT / "results" / "thesis_docx_audit.json"

REQUIRED_HEADINGS = [
    "Abstract",
    "Chapter 1. Introduction",
    "Chapter 2. Literature Review",
    "Chapter 3. Methodology",
    "Chapter 4. Results",
    "Chapter 5. Software, Accessibility and Data Interoperability",
    "Chapter 6. Discussion and Conclusion",
    "Appendix A. Dataset Datasheet Summary",
    "Appendix B. Model Card Summary",
    "Appendix C. Reproducibility Checklist",
    "Appendix D. Defence Demonstration Script",
    "Appendix E. Release Package",
    "Bibliography",
]


def main() -> None:
    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)

    doc = Document(DOCX_PATH)
    headings = [
        paragraph.text.strip()
        for paragraph in doc.paragraphs
        if paragraph.style and paragraph.style.name.startswith("Heading")
    ]
    table_text = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                table_text.append(cell.text)
    text = "\n".join([paragraph.text for paragraph in doc.paragraphs] + table_text)

    with zipfile.ZipFile(DOCX_PATH) as package:
        names = package.namelist()
        media_files = [name for name in names if name.startswith("word/media/")]
        styles_xml = package.read("word/styles.xml").decode("utf-8")
        document_xml = package.read("word/document.xml").decode("utf-8")

    missing_headings = [heading for heading in REQUIRED_HEADINGS if heading not in headings]
    checks = {
        "docx_exists": DOCX_PATH.exists(),
        "required_headings_present": not missing_headings,
        "has_tables": len(doc.tables) >= 8,
        "has_embedded_figures": len(media_files) >= 2,
        "has_results_metrics": "0.8495" in text and "0.8482" in text,
        "has_yolo": "YOLO" in text,
        "has_efficientnetb4": "EfficientNetB4" in text,
        "has_darwin_core": "Darwin Core" in text,
        "has_auslan_caution": "provisional" in text and "AUSLAN" in text,
        "styles_include_heading_blue": "2E74B5" in styles_xml,
        "document_has_page_breaks": 'w:type="page"' in document_xml,
    }

    report = {
        "docx_path": str(DOCX_PATH),
        "paragraphs": len(doc.paragraphs),
        "headings": headings,
        "missing_headings": missing_headings,
        "tables": len(doc.tables),
        "embedded_media_files": media_files,
        "checks": checks,
        "passed": all(checks.values()),
        "note": (
            "Structural DOCX audit only. Visual render QA requires LibreOffice "
            "or another DOCX-to-PDF renderer."
        ),
    }

    RESULTS_PATH.parent.mkdir(parents=True, exist_ok=True)
    RESULTS_PATH.write_text(json.dumps(report, indent=2), encoding="utf-8")

    if not report["passed"]:
        print(json.dumps(report, indent=2))
        raise SystemExit(1)

    print(json.dumps(report, indent=2))


if __name__ == "__main__":
    main()
